"""
Create two versions of the road trip plan:
1. Friend-free markdown version
2. Word docx file for review (with clickable hyperlinks)
"""
import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Read the plan file
plan_path = r"C:\Users\AB Digial\.claude\plans\typed-dancing-adleman.md"
with open(plan_path, 'r', encoding='utf-8') as f:
    content = f.read()

output_dir = r"C:\Users\AB Digial\OneDrive\Documents\Claude"

# ============================================
# TASK 1: Create friend-free version
# ============================================
friend_free = content

friend_free = friend_free.replace(
    " Friend from Greenville, SC will meet you in the Asheville area.", "")
friend_free = friend_free.replace(
    "— weekend day for your friend to join",
    "— weekend day, less crowded at the estate")
friend_free = friend_free.replace(
    "*Visit with friend + lunch*", "*Lunch + explore the area*")
friend_free = friend_free.replace(
    "visit your friend, grab lunch, explore the area. **Allow 1–2 hours.** Near downtown Greenville — walking distance to Main Street, the Commons, and the Swamp Rabbit Trail.",
    "grab lunch, explore the area. **Allow 1–2 hours.** Near downtown Greenville — walking distance to Main Street, the Commons, and the Swamp Rabbit Trail.")
friend_free = friend_free.replace(
    "*Great day to invite your Greenville friend — Hendersonville is only ~45 min from them! Their kids (ages 1 & 4) will love the goats and the children's museum.*\n", "")
friend_free = friend_free.replace(
    " (Perfect for your friend's 1 & 4-year-old too)", "")
friend_free = friend_free.replace(
    " Great for ages 1–8 — perfect for your friend's little ones too.",
    " Great for ages 1–8.")
friend_free = friend_free.replace(
    "*Great day to invite your Greenville friend for a second day! Weekend = easier for them. Their kids (ages 1 & 4) will love the Farmyard animals. Biltmore\'s flat trails work for strollers too.*\n", "")
friend_free = friend_free.replace(
    " Stroller-friendly for your friend's 1-year-old.", "")
friend_free = friend_free.replace(
    " (allow 2+ hours — your son may or may not enjoy this; friend's toddlers probably won't)",
    " (allow 2+ hours — your son may or may not enjoy this)")
friend_free = friend_free.replace(
    "Hendersonville & Flat Rock (friend joins!)",
    "Hendersonville & Flat Rock")
friend_free = friend_free.replace(
    "Biltmore Estate (friend joins!) + Fringe Festival opens",
    "Biltmore Estate + Fringe Festival opens")
friend_free = friend_free.replace(
    "- [ ] **Coordinate with Greenville friend** (2 kids, ages 1 & 4) — best days: **Day 4 Sat, March 14 (Hendersonville)** + **Day 5 Sun, March 15 (Biltmore)**. Both are weekend days — Hendersonville is toddler heaven (goats, children's museum), Biltmore has Farmyard + stroller-friendly trails.\n", "")

friend_free_path = os.path.join(output_dir, "road-trip-plan-no-friends.md")
with open(friend_free_path, 'w', encoding='utf-8') as f:
    f.write(friend_free)
print(f"Friend-free version saved to: {friend_free_path}")

# ============================================
# TASK 2: Create Word docx files with HYPERLINKS
# ============================================

def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')  # 10.5pt
    rPr.append(sz)

    # Font name
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_hyperlink_to_cell(cell, text, url):
    """Add a clickable hyperlink inside a table cell."""
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    add_hyperlink(paragraph, text, url)


def add_run_with_format(paragraph, text, bold=False, italic=False, size=None, color=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    return run


def render_inline(paragraph, text, default_size=None):
    """
    Render inline markdown (bold, italic, links) into a paragraph with proper formatting.
    Handles: **bold**, *italic*, [text](url), and plain text.
    """
    # Pattern to match: **bold**, *italic*, [text](url)
    # Process in order of appearance
    pattern = re.compile(
        r'(\*\*.*?\*\*)'           # bold
        r'|(\*[^*]+?\*)'           # italic
        r'|(\[[^\]]+\]\([^\)]+\))' # link
    )

    pos = 0
    for m in pattern.finditer(text):
        # Add plain text before this match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            if default_size:
                run.font.size = default_size

        if m.group(1):  # Bold
            bold_text = m.group(1)[2:-2]
            # Check if bold text contains links
            link_match = re.search(r'\[([^\]]+)\]\(([^\)]+)\)', bold_text)
            if link_match:
                # Bold text before link
                before = bold_text[:link_match.start()]
                if before:
                    run = paragraph.add_run(before)
                    run.bold = True
                    if default_size:
                        run.font.size = default_size
                # The link itself (bold + hyperlink)
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
                # Bold text after link
                after = bold_text[link_match.end():]
                if after:
                    run = paragraph.add_run(after)
                    run.bold = True
                    if default_size:
                        run.font.size = default_size
            else:
                run = paragraph.add_run(bold_text)
                run.bold = True
                if default_size:
                    run.font.size = default_size

        elif m.group(2):  # Italic
            italic_text = m.group(2)[1:-1]
            run = paragraph.add_run(italic_text)
            run.italic = True
            if default_size:
                run.font.size = default_size

        elif m.group(3):  # Link
            link_match = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', m.group(3))
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))

        pos = m.end()

    # Add remaining plain text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if default_size:
            run.font.size = default_size


def add_table_from_md(doc, table_text):
    """Parse a markdown table and add it to the document with clickable links."""
    lines = [l.strip() for l in table_text.strip().split('\n') if l.strip()]
    if len(lines) < 2:
        return

    # Filter out separator lines
    data_lines = []
    for line in lines:
        stripped = line.strip('|').strip()
        if stripped and not all(c in '-|: ' for c in stripped):
            data_lines.append(line)

    if not data_lines:
        return

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell_text = cell_text.strip()

                # Clear default paragraph
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

                cell_size = Pt(9)

                # Check if cell contains a markdown link
                if re.search(r'\[([^\]]+)\]\(([^\)]+)\)', cell_text):
                    render_inline(p, cell_text, default_size=cell_size)
                else:
                    # Clean markdown bold/italic for plain cells
                    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                    clean = re.sub(r'\*(.*?)\*', r'\1', clean)

                    if i == 0:  # Header row
                        run = p.add_run(clean)
                        run.bold = True
                        run.font.size = cell_size
                    else:
                        # Check for bold in original
                        if '**' in cell_text:
                            render_inline(p, cell_text, default_size=cell_size)
                        else:
                            run = p.add_run(clean)
                            run.font.size = cell_size

    doc.add_paragraph()


def process_paragraph(doc, text, is_bullet=False, indent_level=0):
    """Add a paragraph with inline formatting (bold, italic, links) preserved."""
    # Handle checkbox markers
    text = text.replace('- [x] ', '✅ ').replace('- [ ] ', '⬜ ')

    p = doc.add_paragraph()
    if is_bullet:
        p.style = 'List Bullet'
        if indent_level > 0:
            try:
                p.style = 'List Bullet 2'
            except:
                p.style = 'List Bullet'

    render_inline(p, text)
    return p


def build_docx(md_content, output_path):
    """Build a Word document from markdown content with proper hyperlinks."""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Empty line
        if not stripped:
            if in_table and table_lines:
                add_table_from_md(doc, '\n'.join(table_lines))
                table_lines = []
                in_table = False
            i += 1
            continue

        # Table rows
        if stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table and table_lines:
            add_table_from_md(doc, '\n'.join(table_lines))
            table_lines = []
            in_table = False

        # Horizontal rules
        if stripped in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 80)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(180, 180, 180)
            i += 1
            continue

        # Headers
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            header_text = stripped.lstrip('#').strip()
            # Clean markdown from header but preserve emoji
            header_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', header_text)
            header_clean = re.sub(r'\*(.*?)\*', r'\1', header_clean)
            header_clean = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', header_clean)

            if level == 1:
                h = doc.add_heading(header_clean, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                doc.add_heading(header_clean, level=2)
            elif level == 3:
                doc.add_heading(header_clean, level=3)
            else:
                doc.add_heading(header_clean, level=min(level, 4))
            i += 1
            continue

        # Bullet points (- or *)
        if stripped.startswith('- ') or (stripped.startswith('* ') and not stripped.startswith('**')):
            bullet_text = stripped[2:]
            indent = 0
            # Check for indented sub-bullets
            leading_spaces = len(line) - len(line.lstrip())
            if leading_spaces >= 2:
                indent = 1
            process_paragraph(doc, bullet_text, is_bullet=True, indent_level=indent)
            i += 1
            continue

        # Checkbox items
        if stripped.startswith('- ['):
            process_paragraph(doc, stripped, is_bullet=True)
            i += 1
            continue

        # Regular paragraph
        process_paragraph(doc, stripped)
        i += 1

    # Flush remaining table
    if table_lines:
        add_table_from_md(doc, '\n'.join(table_lines))

    doc.save(output_path)
    print(f"Word document saved to: {output_path}")


# ============================================
# Build both docx files
# ============================================
docx_path = os.path.join(output_dir, "road-trip-plan-v2.docx")
build_docx(content, docx_path)

docx_ff_path = os.path.join(output_dir, "road-trip-plan-no-friends-v2.docx")
build_docx(friend_free, docx_ff_path)

print(f"\n=== ALL FILES CREATED ===")
print(f"1. {friend_free_path}")
print(f"2. {docx_path}")
print(f"3. {docx_ff_path}")
